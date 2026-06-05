import streamlit as st
import os
import time
import warnings
import litellm
from datetime import datetime
from crewai import Agent, Task, Crew, LLM
from crewai_tools import SerperDevTool
from pydantic import BaseModel, Field
from io import BytesIO

# --- NEW FORMATTING IMPORTS ---
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches

# Disable warnings
warnings.filterwarnings("ignore")

# Global Monkey Patch for LiteLLM to prevent structural rejections
original_completion = litellm.completion

def safe_completion(*args, **kwargs):
    if "extra_body" in kwargs and isinstance(kwargs["extra_body"], dict):
        kwargs["extra_body"].pop("cache_breakpoint", None)
        if not kwargs["extra_body"]:
            kwargs.pop("extra_body")
            
    if "messages" in kwargs and isinstance(kwargs["messages"], list):
        for msg in kwargs["messages"]:
            if isinstance(msg, dict):
                msg.pop("cache_breakpoint", None)
                
    kwargs.pop("cache_breakpoint", None)
    return original_completion(*args, **kwargs)

litellm.completion = safe_completion

# --- STRUCTURED OUTCOME SCHEMA ---
class StructuredArticle(BaseModel):
    title: str = Field(description="The catchy, SEO-optimized title of the article.")
    introduction: str = Field(description="The introductory section. Must include an embedded blockquote callout for a core metric or trend.")
    comparative_table_markdown: str = Field(description="A fully formatted Markdown table comparing 3+ core entities with descriptive context hyperlinks embedded inside the cells.")
    body_sections_markdown: str = Field(description="The main editorial content sections separated by clean markdown H3 headers. Each section must have 2-3 detailed paragraphs containing descriptive, context-anchored links.")
    conclusion: str = Field(description="Summarizing wrap-up section and a clear call to action.")


# --- FIXED EXPORT HELPER FUNCTIONS ---

def generate_docx(data: StructuredArticle):
    """Generates a styled Word document cleanly buffered in memory."""
    doc = Document()
    
    # Configure document typography
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title Section
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(data.title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    # Introduction Section
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(data.introduction)
    
    # Comparative Overview Section
    doc.add_heading("Comparative Analysis & Key Resources", level=1)
    doc.add_paragraph(data.comparative_table_markdown)
    
    # Deep-Dive Section
    doc.add_heading("Deep-Dive Content Architecture", level=1)
    doc.add_paragraph(data.body_sections_markdown)
    
    # Conclusion Section
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(data.conclusion)
    
    # Use a secure in-memory bytes stream for Streamlit compatibility
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue() # Returns raw 'bytes'


def generate_pdf(data: StructuredArticle):
    """Generates a clean PDF document, safely cast to standard bytes format."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    def safe_encode(text):
        if not text:
            return ""
        return (
            text.replace("“", '"')
                .replace("”", '"')
                .replace("‘", "'")
                .replace("’", "'")
                .replace("—", "-")
                .replace("–", "-")
                .replace("•", "*")
                .replace("**", "")
                .replace("###", "")
                .replace(">", "")
        )

    # Render Document Title
    pdf.set_font("Helvetica", "B", 18)
    safe_title = safe_encode(data.title)
    pdf.multi_cell(0, 10, safe_title)
    pdf.ln(10)
    
    def add_pdf_section(heading, text):
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, heading, ln=True)
        pdf.ln(2)
        
        pdf.set_font("Helvetica", "", 11)
        clean_text = safe_encode(text)
        pdf.multi_cell(0, 6, clean_text)
        pdf.ln(6)

    add_pdf_section("Introduction", data.introduction)
    add_pdf_section("Comparative Table & Reference Notes", data.comparative_table_markdown)
    add_pdf_section("Deep-Dive Content", data.body_sections_markdown)
    add_pdf_section("Conclusion", data.conclusion)
    
    # CRITICAL FIX: Convert the fpdf2 bytearray explicitly into a standard bytes object
    raw_pdf_data = pdf.output()
    return bytes(raw_pdf_data)


# --- STREAMLIT UI & CONFIGURATIONS ---
st.set_page_config(page_title="Advanced AI Content Crew", page_icon="🚀", layout="centered")

st.title("🚀 Advanced Multi-Agent Content Crew")
st.write("An enhanced system utilizing live web search, parameter tuning, and multi-format downloading.")

# --- SIDEBAR CONTROL PANEL ---
# --- SIDEBAR CONTROL PANEL ---
st.sidebar.header("🔑 API Credentials")

# 1. ALWAYS DECLARE INPUTS FIRST so they never disappear from the UI
groq_api_key = os.environ.get("GROQ_API_KEY") or st.sidebar.text_input(
    "Groq API Key", 
    type="password",
    value=os.environ.get("GROQ_API_KEY", "")
)

serper_api_key = os.environ.get("SERPER_API_KEY") or st.sidebar.text_input(
    "Serper API Key (For Live Search)", 
    type="password",
    value=os.environ.get("SERPER_API_KEY", "")
)

st.sidebar.markdown("---")
st.sidebar.header("🎛️ Generation Settings")

tone = st.sidebar.selectbox(
    "Select Writing Tone",
    ["Professional", "Casual/Conversational", "Technical & Academic", "SEO-Optimized Marketing"]
)

temperature = st.sidebar.slider(
    "LLM Creativity (Temperature)", min_value=0.0, max_value=1.0, value=0.7, step=0.1
)

# 2. RUN THE CONTROLS CHECK AFTER THE SIDEBAR COMPONENT MAP IS FULLY DRAWN
if not groq_api_key:
    st.info("🔑 Please enter your Groq API Key in the sidebar to unlock the generation panel.")
    st.stop()

# --- CACHED CREW EXECUTION FUNCTION ---
@st.cache_data(show_spinner=False)
def run_cached_crew(topic, tone_setting, temp_setting, _groq_key, _serper_key):
    search_tool = None
    if _serper_key:
        os.environ["SERPER_API_KEY"] = _serper_key
        
        # We explicitly enforce the name and description strings to eliminate native fallback triggers
        search_tool = SerperDevTool(
            name="search_the_internet_with_serper",
            description="A precise tool to look up current year news, official corporate documentation, or real-time data trends on Google."
        )

    fast_llm = LLM(model="groq/llama-3.1-8b-instant", api_key=_groq_key, temperature=temp_setting)
    smart_llm = LLM(model="groq/llama-3.3-70b-versatile", api_key=_groq_key, temperature=temp_setting)

    planner = Agent(
        role="Content Research Planner",
        goal=f"Plan factually accurate content architecture on {{topic}} using a {tone_setting} approach.",
        backstory="You extract core structural points, note high-profile sub-concepts, and search for live data.",
        allow_delegation=False, verbose=True, llm=fast_llm,
        tools=[search_tool] if search_tool else []
    )

    writer = Agent(
        role="Content Writer",
        goal=f"Write insightful articles about the topic: {{topic}} matching a {tone_setting} tone perfectly.",
        backstory=f"You write premium copy and maintain a strictly {tone_setting} narrative style.",
        allow_delegation=False, verbose=True, llm=smart_llm
    )

    editor = Agent(
        role="Editor",
        goal="Edit a given blog post to align with professional guidelines and formatting rules.",
        backstory=f"You adjust syntax, pacing, and verify the post strictly embodies a {tone_setting} delivery.",
        allow_delegation=False, verbose=True, llm=smart_llm
    )

    linker = Agent(
        role="Digital Link Optimizer",
        goal="Read a completed text draft and format references into standard Markdown links dynamically.",
        backstory="You convert high-level organizations, cities, or applications mentioned into valid hyperlinks.",
        allow_delegation=False, verbose=True, llm=fast_llm
    )

    # Tasks
    current_year = datetime.now().strftime("%Y")
    plan_description = (
        f"1. Identify trends, structural shifts, and key authoritative entities on {{topic}}.\n"
        f"2. Create a clean article outline highlighting where comparative tables or metric callouts would maximize readability.\n"
        f"3. CRITICAL SEARCH RULE: When searching for data on {{topic}}, target the **current year ({current_year})**.\n"
        f"   Always prioritize official documentation, primary source sites, and corporate/government announcements from {current_year}."
    )

    plan = Task(description=plan_description, expected_output=f"An outline document with verified data notes from {current_year}.", agent=planner)
    
    write = Task(
        description=(
            f"1. Convert the content plan into a premium article on {{topic}} in a {tone_setting} voice.\n"
            f"2. VISUAL READABILITY RULES:\n"
            f"   - Break long prose blocks by introducing a Markdown Table if comparing 2 or more entities.\n"
            f"   - Use Markdown Blockquotes (`>`) to emphasize critical takeaways, warnings, or expert metrics.\n"
            f"   - Ensure each core section header contains exactly 2 to 3 well-paced paragraphs."
        ),
        expected_output="A visually engaging blog post in markdown format with integrated tables and blockquotes.",
        agent=writer
    )

    edit = Task(description="Review and refine the blog post written by the writer.", expected_output="A polished version of the blog post.", agent=editor)

    enrich_links = Task(
        description=(
            "Locate official bodies, software repositories, tools, or major platforms mentioned in the text. "
            "Convert these references into clickable Markdown hyperlinks using descriptive, context-aware action phrases.\n"
            "Organize the gathered facts strictly into the required structured JSON format, ensuring a fully populated markdown "
            "comparative table and an introductory quote block are present."
        ),
        expected_output="A structured database object containing a title, intro, table, body, and conclusion.",
        agent=linker,
        output_pydantic=StructuredArticle
    )

    crew = Crew(agents=[planner, writer, editor, linker], tasks=[plan, write, edit, enrich_links], verbose=True, max_rpm=1)

    result = crew.kickoff(inputs={"topic": topic})
    # Return the raw Pydantic object so we can parse it dynamically into multiple downloads
    return result.pydantic


# --- MAIN INTERFACE ---
# --- MAIN INTERFACE ---
topic = st.text_input("What topic would you like the agents to handle today?", placeholder="e.g., Current Travel Trends Singapore to Tokyo")

if st.button("Launch Crew Execution", type="primary"):
    # 1. RUN THE VALIDATION CHECK *INSIDE* THE BUTTON CLICK EVENT
    if not groq_api_key:
        st.error("🔑 Groq API Key is missing! Please enter your key in the left sidebar to unlock execution.")
    elif not topic.strip():
        st.warning("Please provide a valid topic.")
    else:
        status_box = st.empty()
        max_retries = 3
        
        for attempt in range(max_retries):
            status_box.info(f"🕵️‍♂️ Fetching content for '{topic}' using a {tone} tone... (Checking Cache / Pacing Agents)")
            try:
                # Call the cached helper function which now returns our structured object
                article_data = run_cached_crew(topic, tone, temperature, groq_api_key, serper_api_key)
                
                # Assemble the plain markdown string for immediate on-screen layout visualization
                final_markdown = f"""# {article_data.title}

### Introduction
{article_data.introduction}

### Comparative Overview & Key Resources
{article_data.comparative_table_markdown}

{article_data.body_sections_markdown}

### Conclusion
{article_data.conclusion}
"""
                
                status_box.empty()
                st.success("🎉 Visually optimized article ready!")
                st.subheader("📝 Generated Blog Post")
                
                # Render beautifully to web screen
                st.markdown(final_markdown)
                
                st.markdown("---")
                st.subheader("📥 Export & Download Options")
                
                # Create 3 clean, side-by-side buttons for downloading options
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.download_button(
                        label="📄 Download as Markdown (.md)",
                        data=final_markdown,
                        file_name=f"{topic.lower().replace(' ', '_')}.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
                
                with col2:
                    docx_data = generate_docx(article_data)
                    st.download_button(
                        label="📘 Download as Word (.docx)",
                        data=docx_data,
                        file_name=f"{topic.lower().replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                with col3:
                    pdf_data = generate_pdf(article_data)
                    st.download_button(
                        label="📕 Download as PDF (.pdf)",
                        data=pdf_data,
                        file_name=f"{topic.lower().replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                break
                
            except Exception as e:
                error_msg = str(e)
                if "rate_limit_exceeded" in error_msg.lower() or "ratelimiterror" in error_msg.lower():
                    wait_time = 15.0
                    status_box.warning(f"⏳ Dynamic Rate-limit pause triggered. Resuming in {wait_time:.1f}s...")
                    time.sleep(wait_time)
                else:
                    status_box.error(f"Execution Error: {e}")
                    break